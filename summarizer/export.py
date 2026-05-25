from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime
import io

def generate_txt_report(summary_text: str, sentiment_label: str, sentiment_score: float, source_type: str) -> bytes:
    """Generates a plain text summary report."""
    date_str = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    words_count = len(summary_text.split())
    
    report = f"""==================================================
DOCSNIPPED SUMMARY REPORT
==================================================
Processed on:   {date_str}
Source Type:    {source_type}
Summary Length: {words_count} words
--------------------------------------------------

GENERATED SUMMARY:
{summary_text}

--------------------------------------------------
ANALYSIS INSIGHTS:
Sentiment:      {sentiment_label}
Confidence:     {sentiment_score:.0%}

==================================================
"""
    return report.encode('utf-8')

def generate_docx_report(summary_text: str, sentiment_label: str, sentiment_score: float, source_type: str) -> bytes:
    """Generates a professionally styled Microsoft Word document summary report."""
    doc = docx.Document()
    
    # Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styling colors
    color_primary = RGBColor(124, 111, 205)  # #7C6FCD
    color_dark = RGBColor(30, 27, 75)        # #1e1b4b
    color_muted = RGBColor(75, 85, 99)       # #4b5563
    
    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("DocSnipped Summary Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = color_dark
    title_p.paragraph_format.space_after = Pt(2)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("AI-Powered Document Summarization & Sentiment Analysis")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = color_muted
    sub_p.paragraph_format.space_after = Pt(18)
    
    # Divider line
    p_line = doc.add_paragraph()
    p_line_run = p_line.add_run("─" * 60)
    p_line_run.font.color.rgb = color_primary
    p_line.paragraph_format.space_after = Pt(12)
    
    # Metadata
    date_str = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    words_count = len(summary_text.split())
    
    meta_p = doc.add_paragraph()
    
    run_s = meta_p.add_run("Source Type: ")
    run_s.bold = True
    meta_p.add_run(f"{source_type}    |    ")
    
    run_d = meta_p.add_run("Processed on: ")
    run_d.bold = True
    meta_p.add_run(f"{date_str}    |    ")
    
    run_w = meta_p.add_run("Summary Length: ")
    run_w.bold = True
    meta_p.add_run(f"{words_count} words")
    
    for run in meta_p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = color_muted
        
    meta_p.paragraph_format.space_after = Pt(18)
    
    # Summary Section Heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("Generated Summary")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = color_primary
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    body = doc.add_paragraph()
    body_run = body.add_run(summary_text)
    body_run.font.name = 'Arial'
    body_run.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.25
    body.paragraph_format.space_after = Pt(18)
    
    # Sentiment Section Heading
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Analysis Insights")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = color_primary
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    sent_p = doc.add_paragraph()
    
    run_sent_lbl = sent_p.add_run("SENTIMENT: ")
    run_sent_lbl.font.name = 'Arial'
    run_sent_lbl.font.size = Pt(10)
    run_sent_lbl.font.bold = True
    run_sent_lbl.font.color.rgb = color_primary
    
    run_sent_val = sent_p.add_run(sentiment_label)
    run_sent_val.font.name = 'Arial'
    run_sent_val.font.size = Pt(14)
    run_sent_val.font.bold = True
    if sentiment_label == "POSITIVE":
        run_sent_val.font.color.rgb = RGBColor(22, 163, 74)
    elif sentiment_label == "NEGATIVE":
        run_sent_val.font.color.rgb = RGBColor(220, 38, 38)
    else:
        run_sent_val.font.color.rgb = color_muted
        
    sent_p.add_run(f"\nConfidence score: {sentiment_score:.0%}")
    sent_p.runs[-1].font.name = 'Arial'
    sent_p.runs[-1].font.size = Pt(9.5)
    sent_p.runs[-1].font.color.rgb = color_muted
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes

def generate_pdf_report(summary_text: str, sentiment_label: str, sentiment_score: float, source_type: str) -> bytes:
    """Generates a beautifully designed, high-fidelity PDF summary report using reportlab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom color palette matching DocSnipped branding
    brand_primary = colors.HexColor("#7C6FCD")
    brand_dark = colors.HexColor("#1e1b4b")
    text_dark = colors.HexColor("#111827")
    muted_grey = colors.HexColor("#4b5563")
    border_color = colors.HexColor("#e2e8f0")
    
    # Custom styles definitions
    title_style = ParagraphStyle(
        'DocSnippedTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=brand_dark,
        spaceAfter=6
    )
    
    subtitle_style = ParagraphStyle(
        'DocSnippedSub',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=muted_grey,
        spaceAfter=20
    )
    
    section_heading = ParagraphStyle(
        'DocSnippedSecHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=brand_primary,
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'DocSnippedBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=16,
        textColor=text_dark,
        spaceAfter=12
    )
    
    meta_style = ParagraphStyle(
        'DocSnippedMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=muted_grey
    )
    
    sentiment_title_style = ParagraphStyle(
        'DocSnippedSentTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=brand_primary
    )
    
    sentiment_val_style = ParagraphStyle(
        'DocSnippedSentVal',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#16a34a") if sentiment_label == "POSITIVE" else colors.HexColor("#dc2626") if sentiment_label == "NEGATIVE" else muted_grey
    )
    
    story = []
    
    # Title & Subtitle
    story.append(Paragraph("DocSnipped Summary Report", title_style))
    story.append(Paragraph("AI-Powered Document Summarization & Sentiment Analysis", subtitle_style))
    
    # Colored visual anchor bar
    d_bar = Table([[""]], colWidths=[doc.width], rowHeights=[3])
    d_bar.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), brand_primary),
        ('PADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(d_bar)
    story.append(Spacer(1, 15))
    
    # Metadata block
    date_str = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    words_count = len(summary_text.split())
    
    meta_data = [
        [Paragraph(f"<b>Source Type:</b> {source_type}", meta_style), 
         Paragraph(f"<b>Processed on:</b> {date_str}", meta_style)],
        [Paragraph(f"<b>Summary Length:</b> {words_count} words", meta_style), ""]
    ]
    meta_table = Table(meta_data, colWidths=[doc.width/2.0, doc.width/2.0])
    meta_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 15))
    
    # Summary Content Section
    story.append(Paragraph("Generated Summary", section_heading))
    # Replace newlines with html breaks for PDF format
    formatted_summary = summary_text.replace("\n", "<br/>")
    story.append(Paragraph(formatted_summary, body_style))
    story.append(Spacer(1, 10))
    
    # Insights Section (styled box container)
    story.append(Paragraph("Analysis Insights", section_heading))
    
    sent_data = [
        [Paragraph("SENTIMENT", sentiment_title_style)],
        [Paragraph(sentiment_label, sentiment_val_style)],
        [Paragraph(f"Confidence score: {sentiment_score:.0%}", meta_style)]
    ]
    sent_table = Table(sent_data, colWidths=[200])
    sent_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8f9fc")),
        ('BOX', (0,0), (-1,-1), 1, border_color),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
    ]))
    
    story.append(sent_table)
    
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes
